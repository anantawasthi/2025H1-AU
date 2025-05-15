import pandas as pd
import numpy as np
from collections import Counter
from docx import Document

# Load data
file_path = "../synthetic_classification_dataset.csv"
df = pd.read_csv(file_path)

# Data overview
n_rows, n_cols = df.shape
col_names = df.columns.tolist()
dtypes = df.dtypes.astype(str).to_dict()
missing_count = df.isnull().sum().to_dict()
missing_pct = (df.isnull().mean() * 100).round(2).to_dict()
unique_vals = {col: df[col].nunique() for col in df.columns}
describe = df.describe(include='all').T

# Create DOCX report
doc = Document()
doc.add_heading('Data Overview Report', 0)
doc.add_paragraph(f"Number of rows: {n_rows}")
doc.add_paragraph(f"Number of columns: {n_cols}")
doc.add_heading('Column Details', level=1)

table = doc.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Column Name'
hdr_cells[1].text = 'Data Type'
hdr_cells[2].text = 'Missing Values (Count)'
hdr_cells[3].text = 'Missing Values (%)'
hdr_cells[4].text = 'Unique Values'
hdr_cells[5].text = 'Sample/Stats'

for col in col_names:
    row_cells = table.add_row().cells
    row_cells[0].text = col
    row_cells[1].text = dtypes[col]
    row_cells[2].text = str(missing_count[col])
    row_cells[3].text = str(missing_pct[col])
    row_cells[4].text = str(unique_vals[col])
    if col in describe.index:
        stats = describe.loc[col]
        if dtypes[col] in ['int64', 'float64']:
            stat_str = f"mean={stats['mean']:.2f}, std={stats['std']:.2f}, min={stats['min']}, max={stats['max']}"
        else:
            stat_str = f"top={stats['top']}, freq={stats['freq']}"
    else:
        stat_str = str(df[col].unique()[:3])
    row_cells[5].text = stat_str

doc.save("eda/data_overview.docx")
print("DOCX report saved as eda/data_overview.docx")
