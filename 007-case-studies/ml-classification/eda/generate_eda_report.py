import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import os
from docx import Document

# Load data
file_path = '../synthetic_classification_dataset.csv'
df = pd.read_csv(file_path)

# Create output directory for plots
os.makedirs('eda/plots', exist_ok=True)

# Univariate analysis
univariate_insights = []
for col in df.columns:
    plt.figure(figsize=(6,4))
    if df[col].dtype in ['int64', 'float64'] and col != 'Customer_ID':
        sns.histplot(df[col].dropna(), kde=True)
        plt.title(f'Histogram of {col}')
        plt.savefig(f'eda/plots/hist_{col}.png')
        plt.close()
        desc = df[col].describe()
        univariate_insights.append(f"{col}: mean={desc['mean']:.2f}, std={desc['std']:.2f}, min={desc['min']}, max={desc['max']}")
    elif col != 'Customer_ID':
        sns.countplot(x=col, data=df)
        plt.title(f'Bar plot of {col}')
        plt.savefig(f'eda/plots/bar_{col}.png')
        plt.close()
        counts = df[col].value_counts()
        univariate_insights.append(f"{col}:\n{counts.to_string()}")

# Bivariate analysis: Churn vs other features
bivariate_insights = []
target = 'Customer_Churn'
for col in df.columns:
    if col == target or col == 'Customer_ID':
        continue
    plt.figure(figsize=(6,4))
    if df[col].dtype in ['int64', 'float64']:
        sns.boxplot(x=target, y=col, data=df)
        plt.title(f'{col} by {target}')
        plt.savefig(f'eda/plots/box_{col}_by_{target}.png')
        plt.close()
        means = df.groupby(target)[col].mean()
        bivariate_insights.append(f"{col} mean by {target}:\n{means.to_string()}")
    else:
        ct = pd.crosstab(df[col], df[target])
        ct.plot(kind='bar', stacked=True)
        plt.title(f'{col} vs {target}')
        plt.savefig(f'eda/plots/bar_{col}_by_{target}.png')
        plt.close()
        bivariate_insights.append(f"{col} vs {target}:\n{ct.to_string()}")

# Skewed, imbalanced, or non-informative fields
skewed = []
for col in df.select_dtypes(include=['int64', 'float64']).columns:
    if abs(df[col].skew()) > 1:
        skewed.append(col)
imbalanced = []
churn_counts = df[target].value_counts(normalize=True)
if churn_counts.min() < 0.2:
    imbalanced.append(target)
non_informative = [col for col in df.columns if df[col].nunique() <= 1]

# Create DOCX report
doc = Document()
doc.add_heading('EDA Report', 0)
doc.add_heading('Univariate Analysis', level=1)
for insight in univariate_insights:
    doc.add_paragraph(insight)
doc.add_heading('Bivariate Analysis', level=1)
for insight in bivariate_insights:
    doc.add_paragraph(insight)
doc.add_heading('Skewed Fields', level=1)
doc.add_paragraph(", ".join(skewed) if skewed else "None")
doc.add_heading('Imbalanced Fields', level=1)
doc.add_paragraph(", ".join(imbalanced) if imbalanced else "None")
doc.add_heading('Non-informative Fields', level=1)
doc.add_paragraph(", ".join(non_informative) if non_informative else "None")
doc.save('eda/eda_report.docx')
print('EDA DOCX report and plots saved.')
