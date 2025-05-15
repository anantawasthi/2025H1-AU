import pandas as pd
from xgboost import XGBClassifier
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, roc_auc_score, confusion_matrix, roc_curve, classification_report
import matplotlib.pyplot as plt
import joblib
from docx import Document

# Load data
train = pd.read_csv('model-data/train.csv')
test = pd.read_csv('model-data/test.csv')

# Prepare features and target
X_train = train.drop(['Customer_ID', 'Customer_Churn'], axis=1)
y_train = train['Customer_Churn']
X_test = test.drop(['Customer_ID', 'Customer_Churn'], axis=1)
y_test = test['Customer_Churn']

# Train XGBoost
xgb = XGBClassifier(use_label_encoder=False, eval_metric='logloss', random_state=42)
xgb.fit(X_train, y_train)

# Predict and evaluate
y_pred = xgb.predict(X_test)
y_prob = xgb.predict_proba(X_test)[:, 1]
acc = accuracy_score(y_test, y_pred)
prec = precision_score(y_test, y_pred)
rec = recall_score(y_test, y_pred)
f1 = f1_score(y_test, y_pred)
auc = roc_auc_score(y_test, y_prob)
cm = confusion_matrix(y_test, y_pred)
report = classification_report(y_test, y_pred)

# ROC Curve
fpr, tpr, _ = roc_curve(y_test, y_prob)
plt.figure()
plt.plot(fpr, tpr, label=f'ROC curve (AUC = {auc:.2f})')
plt.plot([0, 1], [0, 1], 'k--')
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('ROC Curve - XGBoost')
plt.legend(loc='lower right')
plt.savefig('model-performance/xgboost_roc_curve.png')
plt.close()

# Save model
joblib.dump(xgb, 'model/xgboost_model.pkl')

# Save report
doc = Document()
doc.add_heading('XGBoost Model Report', 0)
doc.add_heading('Performance Metrics', level=1)
doc.add_paragraph(f'Accuracy: {acc:.3f}')
doc.add_paragraph(f'Precision: {prec:.3f}')
doc.add_paragraph(f'Recall: {rec:.3f}')
doc.add_paragraph(f'F1-Score: {f1:.3f}')
doc.add_paragraph(f'ROC-AUC: {auc:.3f}')
doc.add_heading('Confusion Matrix', level=1)
doc.add_paragraph(str(cm))
doc.add_heading('Classification Report', level=1)
doc.add_paragraph(report)
doc.add_heading('Feature Importances', level=1)
for name, score in zip(X_train.columns, xgb.feature_importances_):
    doc.add_paragraph(f'{name}: {score:.4f}')
doc.add_paragraph('See model-performance/xgboost_roc_curve.png for ROC curve.')
doc.save('model-performance/xgboost_report.docx')
print('XGBoost model and report saved.')
