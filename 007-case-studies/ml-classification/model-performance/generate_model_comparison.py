import pandas as pd
import joblib
from docx import Document

# Load reports
logreg_doc = 'model-performance/logistic_regression_report.docx'
xgb_doc = 'model-performance/xgboost_report.docx'

# Read metrics from both reports (manual summary for now)
logreg_metrics = {
    'Accuracy': 'see report',
    'F1-Score': 'see report',
    'AUC': 'see report',
    'Strengths': 'Simple, interpretable, fast, good baseline',
    'Weaknesses': 'May underfit, less powerful for complex patterns',
    'Fairness': 'Check bias in confusion matrix and metrics by group',
    'Explainability': 'High (coefficients)'
}
xgb_metrics = {
    'Accuracy': 'see report',
    'F1-Score': 'see report',
    'AUC': 'see report',
    'Strengths': 'Handles nonlinearity, high accuracy, robust to outliers',
    'Weaknesses': 'Less interpretable, can overfit, slower',
    'Fairness': 'Check bias in confusion matrix and metrics by group',
    'Explainability': 'Medium (feature importance, SHAP possible)'
}

# Recommendation (example)
recommendation = 'If accuracy and AUC are significantly higher for XGBoost, use XGBoost. If interpretability is critical, use Logistic Regression.'

# Save summary
doc = Document()
doc.add_heading('Model Comparison Summary', 0)
doc.add_heading('Logistic Regression', level=1)
for k, v in logreg_metrics.items():
    doc.add_paragraph(f'{k}: {v}')
doc.add_heading('XGBoost', level=1)
for k, v in xgb_metrics.items():
    doc.add_paragraph(f'{k}: {v}')
doc.add_heading('Recommendation', level=1)
doc.add_paragraph(recommendation)
doc.save('model-performance/model_comparison_summary.docx')
print('Model comparison summary saved.')
