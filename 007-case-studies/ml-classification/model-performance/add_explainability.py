import pandas as pd
import joblib
import shap
import matplotlib.pyplot as plt
from docx import Document

# Load model and data
model = joblib.load('../model/final_model.pkl')
train = pd.read_csv('../model-data/train.csv')
test = pd.read_csv('../model-data/test.csv')

# Prepare features
X_train = train.drop(['Customer_ID', 'Customer_Churn'], axis=1)
X_test = test.drop(['Customer_ID', 'Customer_Churn'], axis=1)

# SHAP explainability
explainer = shap.Explainer(model, X_train)
shap_values = explainer(X_test)

# Summary plot
plt.figure()
shap.summary_plot(shap_values, X_test, show=False)
plt.tight_layout()
plt.savefig('../model-performance/shap_summary_plot.png')
plt.close()

# Feature importance (mean(|SHAP|))
shap_importance = pd.DataFrame({
    'feature': X_test.columns,
    'mean_abs_shap': np.abs(shap_values.values).mean(axis=0)
}).sort_values('mean_abs_shap', ascending=False)
shap_importance.to_csv('../model-performance/shap_feature_importance.csv', index=False)

# Add to model comparison docx
doc = Document('../model-performance/model_comparison_summary.docx')
doc.add_heading('Explainability (SHAP)', level=1)
doc.add_paragraph('SHAP summary plot saved as shap_summary_plot.png.')
doc.add_paragraph('Top features by mean(|SHAP|):')
for i, row in shap_importance.head(10).iterrows():
    doc.add_paragraph(f"{row['feature']}: {row['mean_abs_shap']:.4f}")
doc.save('../model-performance/model_comparison_summary.docx')
print('SHAP explainability added to model comparison summary.')
