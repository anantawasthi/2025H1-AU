import pandas as pd
import numpy as np
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler, OneHotEncoder, LabelEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
import joblib
from docx import Document

# Load data
file_path = "../synthetic_classification_dataset.csv"
try:
    df = pd.read_csv(file_path)
except FileNotFoundError:
    df = pd.read_csv("synthetic_classification_dataset.csv")

# Identify columns
target = 'Customer_Churn'
id_col = 'Customer_ID'
num_cols = ['Age', 'Monthly_Spend', 'Tenure_Years', 'Support_Tickets_Raised', 'City_Tier']
cat_cols = ['Gender', 'Education_Level', 'Used_App_Last_Month']

# Impute missing values (if any)
num_imputer = SimpleImputer(strategy='median')
cat_imputer = SimpleImputer(strategy='most_frequent')
df[num_cols] = num_imputer.fit_transform(df[num_cols])
df[cat_cols] = cat_imputer.fit_transform(df[cat_cols])

# Outlier treatment (clip to 1st/99th percentile)
for col in num_cols:
    lower, upper = np.percentile(df[col], [1, 99])
    df[col] = np.clip(df[col], lower, upper)

# Encode categorical features
# Use OneHot for low cardinality, Label for binary
encoded_features = []
encoders = {}
for col in cat_cols:
    if df[col].nunique() == 2:
        le = LabelEncoder()
        df[col] = le.fit_transform(df[col])
        encoders[col] = le
        encoded_features.append(col)
    else:
        ohe = OneHotEncoder(sparse_output=False, drop='first')
        ohe_arr = ohe.fit_transform(df[[col]])
        ohe_cols = [f"{col}_{cat}" for cat in ohe.categories_[0][1:]]
        df[ohe_cols] = ohe_arr
        encoders[col] = ohe
        encoded_features.extend(ohe_cols)
        df.drop(columns=[col], inplace=True)

# Scale numerical features
scaler = StandardScaler()
df[num_cols] = scaler.fit_transform(df[num_cols])

# New feature: Spend per year of tenure (if tenure > 0)
df['Spend_per_Year'] = df['Monthly_Spend'] / (df['Tenure_Years'] + 1)

# Save transformed dataset summary
doc = Document()
doc.add_heading('Feature Engineering Report', 0)
doc.add_paragraph('Missing values imputed (median for numerics, mode for categoricals).')
doc.add_paragraph('Outliers clipped to 1st and 99th percentiles for numerics.')
doc.add_paragraph('Categorical features encoded (Label/OneHot based on cardinality).')
doc.add_paragraph('Numerical features scaled (StandardScaler).')
doc.add_paragraph('New feature created: Spend_per_Year.')
doc.add_heading('Transformed Data Sample', level=1)
doc.add_paragraph(str(df.head(5)))
doc.save('feature-engineering/feature_engineering.docx')
print('Feature engineering DOCX report saved.')

# Save processed data for next steps
df.to_csv('feature-engineering/processed_data.csv', index=False)
# Save encoders and scaler for scoring script
joblib.dump({'scaler': scaler, 'encoders': encoders, 'num_cols': num_cols, 'cat_cols': cat_cols}, 'model/fe_pipeline.joblib')
